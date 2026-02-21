from __future__ import annotations

import pytest

# ── Text Extraction Tests ──────────────────────────────────────────────────────


class TestTextCleaning:
    def test_clean_text_normalizes_whitespace(self):
        from app.services.extractor import clean_text

        raw = "Hello   World\n\n\n\nFoo"
        result = clean_text(raw)
        assert "   " not in result
        assert result.count("\n\n\n") == 0

    def test_clean_text_removes_page_numbers(self):
        from app.services.extractor import clean_text

        raw = "Some content\n\n  3  \n\nMore content"
        result = clean_text(raw)
        assert "  3  " not in result

    def test_clean_text_normalizes_bullets(self):
        from app.services.extractor import clean_text

        raw = "• Python\n• JavaScript\n* Go"
        result = clean_text(raw)
        assert "•" not in result

    def test_clean_text_preserves_content(self):
        from app.services.extractor import clean_text

        raw = "Python Developer with 5 years of experience in Django and FastAPI"
        result = clean_text(raw)
        assert "Python Developer" in result
        assert "Django" in result
        assert "FastAPI" in result

    def test_clean_text_handles_empty(self):
        from app.services.extractor import clean_text

        assert clean_text("") == ""
        assert clean_text("   \n\n   ") == ""


class TestDocxExtraction:
    def test_extract_docx(self, tmp_path):
        """Create a minimal DOCX and test extraction."""
        try:
            from docx import Document

            from app.services.extractor import extract_text_from_docx

            doc = Document()
            doc.add_paragraph("John Doe")
            doc.add_paragraph("Python Developer")
            doc.add_paragraph("Skills: Python, FastAPI, PostgreSQL")

            docx_path = tmp_path / "test_resume.docx"
            doc.save(str(docx_path))

            text = extract_text_from_docx(str(docx_path))
            assert "Python Developer" in text
            assert "FastAPI" in text
        except ImportError:
            pytest.skip("python-docx not installed")

    def test_extract_docx_with_table(self, tmp_path):
        """Test DOCX extraction includes table content."""
        try:
            from docx import Document

            from app.services.extractor import extract_text_from_docx

            doc = Document()
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Skill"
            table.cell(0, 1).text = "Level"
            table.cell(1, 0).text = "Python"
            table.cell(1, 1).text = "Expert"

            docx_path = tmp_path / "test_table.docx"
            doc.save(str(docx_path))

            text = extract_text_from_docx(str(docx_path))
            assert "Python" in text
        except ImportError:
            pytest.skip("python-docx not installed")


# ── Skills Extraction Tests ────────────────────────────────────────────────────


class TestSkillsExtraction:
    SAMPLE_SKILLS = [
        "python",
        "javascript",
        "fastapi",
        "docker",
        "postgresql",
        "machine learning",
        "react",
        "aws",
        "kubernetes",
        "sql",
    ]

    def test_extract_known_skills(self):
        from app.services.nlp_pipeline import extract_skills_from_text

        text = "I have experience with Python, FastAPI, and Docker."
        skills = extract_skills_from_text(text, self.SAMPLE_SKILLS)
        assert "python" in skills
        assert "fastapi" in skills
        assert "docker" in skills

    def test_extract_case_insensitive(self):
        from app.services.nlp_pipeline import extract_skills_from_text

        text = "PYTHON developer with JAVASCRIPT experience"
        skills = extract_skills_from_text(text, self.SAMPLE_SKILLS)
        assert "python" in skills
        assert "javascript" in skills

    def test_extract_no_false_positives(self):
        from app.services.nlp_pipeline import extract_skills_from_text

        # "sql" should not match "nosql" as a word boundary
        text = "I work with NoSQL databases"
        skills = extract_skills_from_text(text, ["sql"])
        assert "sql" not in skills

    def test_extract_multi_word_skills(self):
        from app.services.nlp_pipeline import extract_skills_from_text

        text = "Experience with machine learning and natural language processing"
        skills = extract_skills_from_text(text, ["machine learning", "natural language processing"])
        assert "machine learning" in skills
        assert "natural language processing" in skills

    def test_extract_empty_text(self):
        from app.services.nlp_pipeline import extract_skills_from_text

        skills = extract_skills_from_text("", self.SAMPLE_SKILLS)
        assert skills == []

    def test_skill_overlap(self):
        from app.services.nlp_pipeline import compute_skill_overlap

        resume_skills = ["python", "docker", "fastapi", "react"]
        jd_skills = ["python", "docker", "kubernetes", "aws"]
        matching, missing = compute_skill_overlap(resume_skills, jd_skills)
        assert "python" in matching
        assert "docker" in matching
        assert "kubernetes" in missing
        assert "aws" in missing
        assert "fastapi" not in missing  # not in JD

    def test_skill_overlap_no_match(self):
        from app.services.nlp_pipeline import compute_skill_overlap

        matching, missing = compute_skill_overlap(["python"], ["java"])
        assert matching == []
        assert "java" in missing


# ── Section Detection Tests ────────────────────────────────────────────────────


class TestSectionDetection:
    def test_detect_experience_section(self):
        from app.services.nlp_pipeline import detect_sections

        text = "John Doe\n\nExperience\nSoftware Engineer at Acme Corp 2020-2023\n\nEducation\nBS Computer Science"
        sections = detect_sections(text)
        assert "experience" in sections
        assert "education" in sections

    def test_detect_skills_section(self):
        from app.services.nlp_pipeline import detect_sections

        text = "Name\n\nSkills\nPython, JavaScript, Docker\n\nProjects\nBuilt a web app"
        sections = detect_sections(text)
        assert "skills" in sections
        assert "projects" in sections

    def test_fallback_to_header(self):
        from app.services.nlp_pipeline import detect_sections

        text = "Some random text without any section headings"
        sections = detect_sections(text)
        assert "header" in sections

    def test_certifications_section(self):
        from app.services.nlp_pipeline import detect_sections

        text = "Work\n\nCertifications\nAWS Certified Solutions Architect"
        sections = detect_sections(text)
        assert "certifications" in sections


# ── Similarity Scoring Tests ───────────────────────────────────────────────────


class TestSimilarityScoring:
    def test_cosine_similarity_identical(self):
        from app.services.embedder import cosine_similarity

        vec = [1.0, 0.0, 0.0]
        assert abs(cosine_similarity(vec, vec) - 1.0) < 1e-6

    def test_cosine_similarity_orthogonal(self):
        from app.services.embedder import cosine_similarity

        a = [1.0, 0.0, 0.0]
        b = [0.0, 1.0, 0.0]
        assert abs(cosine_similarity(a, b)) < 1e-6

    def test_cosine_similarity_opposite(self):
        from app.services.embedder import cosine_similarity

        a = [1.0, 0.0]
        b = [-1.0, 0.0]
        assert abs(cosine_similarity(a, b) - (-1.0)) < 1e-6

    def test_cosine_similarity_zero_vector(self):
        from app.services.embedder import cosine_similarity

        assert cosine_similarity([0.0, 0.0], [1.0, 0.0]) == 0.0

    def test_score_mapping_max(self):
        from app.services.embedder import similarity_to_score

        assert similarity_to_score(1.0) == 100.0

    def test_score_mapping_min(self):
        from app.services.embedder import similarity_to_score

        assert similarity_to_score(-1.0) == 0.0

    def test_score_mapping_neutral(self):
        from app.services.embedder import similarity_to_score

        assert similarity_to_score(0.0) == 50.0

    def test_score_clamped(self):
        from app.services.embedder import similarity_to_score

        assert similarity_to_score(2.0) == 100.0
        assert similarity_to_score(-2.0) == 0.0

    def test_score_deterministic(self):
        """Same inputs must always produce same output."""
        from app.services.embedder import similarity_to_score

        for val in [-0.5, 0.0, 0.3, 0.7, 1.0]:
            assert similarity_to_score(val) == similarity_to_score(val)


# ── Redaction Tests ────────────────────────────────────────────────────────────


class TestRedaction:
    def test_redact_email(self):
        from app.services.redactor import redact_text

        text = "Contact me at john.doe@example.com for more info"
        result = redact_text(text)
        assert "john.doe@example.com" not in result
        assert "[EMAIL]" in result

    def test_redact_phone(self):
        from app.services.redactor import redact_text

        text = "Call me at +1 555-123-4567"
        result = redact_text(text)
        assert "555-123-4567" not in result

    def test_preserve_skills(self):
        from app.services.redactor import redact_text

        text = "Skills: Python, FastAPI, Docker, PostgreSQL"
        result = redact_text(text)
        assert "Python" in result
        assert "FastAPI" in result
        assert "Docker" in result

    def test_redact_empty(self):
        from app.services.redactor import redact_text

        assert redact_text("") == ""
