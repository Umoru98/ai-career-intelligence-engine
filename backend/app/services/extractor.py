from __future__ import annotations

import re


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file using pdfplumber (preferred) or pypdf fallback."""
    text_parts: list[str] = []

    try:
        import pdfplumber

        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
    except Exception as pdf_err:
        # Fallback to pypdf
        try:
            from pypdf import PdfReader

            reader = PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        except Exception as pypdf_err:
            raise RuntimeError(f"PDF extraction failed. pdfplumber: {pdf_err}; pypdf: {pypdf_err}")

    raw = "\n".join(text_parts)
    if not raw.strip():
        raise RuntimeError(
            "Text not extractable: PDF appears to be scanned or image-based. "
            "OCR is not enabled. Please provide a text-selectable PDF."
        )
    return raw


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document

        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        raw = "\n".join(paragraphs)
        if not raw.strip():
            raise RuntimeError("DOCX file appears to be empty or contains no extractable text.")
        return raw
    except Exception as e:
        raise RuntimeError(f"DOCX extraction failed: {e}")


def extract_text(file_path: str, content_type: str) -> str:
    """Dispatch extraction based on content type."""
    if content_type == "application/pdf":
        return extract_text_from_pdf(file_path)
    elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported content type: {content_type}")


# ── Text Cleaning ──────────────────────────────────────────────────────────────

_REPEATED_WHITESPACE = re.compile(r"[ \t]+")
_REPEATED_NEWLINES = re.compile(r"\n{3,}")
_PAGE_NUMBERS = re.compile(r"^\s*\d+\s*$", re.MULTILINE)
_BULLET_NORMALIZATION = re.compile(r"^[\u2022\u2023\u25E6\u2043\u2219\-\*]\s+", re.MULTILINE)


def clean_text(raw_text: str) -> str:
    """
    Normalize whitespace, remove page numbers, normalize bullets.
    Does NOT remove content — only cleans formatting artifacts.
    """
    text = raw_text

    # Normalize unicode whitespace
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Remove standalone page numbers
    text = _PAGE_NUMBERS.sub("", text)

    # Normalize bullet characters to dash
    text = _BULLET_NORMALIZATION.sub("- ", text)

    # Collapse repeated horizontal whitespace
    text = _REPEATED_WHITESPACE.sub(" ", text)

    # Collapse excessive blank lines (keep max 2)
    text = _REPEATED_NEWLINES.sub("\n\n", text)

    return text.strip()
